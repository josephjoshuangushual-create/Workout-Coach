from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand palette
NAVY = RGBColor(0x0F, 0x2A, 0x47)
TEAL = RGBColor(0x16, 0x8A, 0x8A)
SLATE = RGBColor(0x3A, 0x4A, 0x5A)
LIGHT = RGBColor(0x6B, 0x7A, 0x8A)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = SLATE
pf = normal.paragraph_format
pf.space_after = Pt(8)
pf.line_spacing = 1.18

# Margins
for s in doc.sections:
    s.top_margin = Inches(0.9)
    s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)


def shade(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def bottom_border(paragraph, color_hex, size=12):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_runs(paragraph, segments, base_color=SLATE):
    # segments: list of (text, bold)
    for text, bold in segments:
        r = paragraph.add_run(text)
        r.bold = bold
        r.font.color.rgb = base_color
    return paragraph


def parse_bold(text):
    # split **bold** markers
    parts = []
    i = 0
    while i < len(text):
        start = text.find("**", i)
        if start == -1:
            parts.append((text[i:], False))
            break
        if start > i:
            parts.append((text[i:start], False))
        end = text.find("**", start + 2)
        if end == -1:
            parts.append((text[start:], False))
            break
        parts.append((text[start + 2:end], True))
        i = end + 2
    return parts


def body(text, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    add_runs(p, parse_bold(text))
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    add_runs(p, parse_bold(text))
    return p


def subhead(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = TEAL
    return p


def question_block(qnum, qtext):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    shade(p, "0F2A47")
    p.paragraph_format.left_indent = Inches(0.08)
    rlabel = p.add_run("  " + qnum + "   ")
    rlabel.bold = True
    rlabel.font.size = Pt(12)
    rlabel.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # question text on next paragraph
    q = doc.add_paragraph()
    q.paragraph_format.space_before = Pt(6)
    q.paragraph_format.space_after = Pt(10)
    rq = q.add_run(qtext)
    rq.bold = True
    rq.font.size = Pt(12.5)
    rq.font.color.rgb = NAVY
    bottom_border(q, "168A8A", size=10)


# ---------------- COVER / HEADER ----------------
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(6)
title.paragraph_format.space_after = Pt(0)
rt = title.add_run("SEO Content Screening Questionnaire")
rt.bold = True
rt.font.size = Pt(22)
rt.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(2)
rs = sub.add_run("Candidate Response")
rs.font.size = Pt(13)
rs.font.color.rgb = TEAL
rs.bold = True

meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(4)
rm = meta.add_run("Joseph  |  SEO Content & Strategy  |  May 2026")
rm.font.size = Pt(10.5)
rm.font.color.rgb = LIGHT
bottom_border(meta, "0F2A47", size=18)

intro = doc.add_paragraph()
intro.paragraph_format.space_before = Pt(8)
intro.paragraph_format.space_after = Pt(4)
ri = intro.add_run(
    "Prepared in response to the screening questionnaire covering ownership of large-scale "
    "SEO content initiatives and the use of AI within a repeatable, quality-controlled content "
    "production process."
)
ri.font.size = Pt(10.5)
ri.italic = True
ri.font.color.rgb = LIGHT

# ---------------- Q1 ----------------
question_block(
    "Q1",
    "Describe the largest SEO content initiative you personally owned from brief to publication. "
    "How many pages did you take live, and over what timeframe?",
)

body(
    "The largest content initiative I personally owned was the complete content rebuild and expansion of "
    "**Med Virtual**, a US-based medical virtual staffing agency headquartered in California that serves "
    "clients across the entire United States. I owned it end to end, from the initial audit and ideation "
    "through site architecture, keyword research, content briefs, and copywriting, then handing finished "
    "briefs to UI/UX for design and on into development. I also continued to own the pages after launch, "
    "updating them to stay aligned with Google's evolving best practices and algorithm updates."
)

subhead("The starting point and the problem I was solving")
body(
    "When I began, the site was essentially undeveloped: roughly **12 pages** of thin, generic content on a "
    "domain only two to three months old, with effectively no domain authority. When I ran the audit, the "
    "pages were not up to par with our competitors, and, critically, they did not reflect the actual roles or "
    "services the company offered. The business served US clients only, so I needed an architecture engineered "
    "to capture **US traffic specifically**. The ICP was doctors and private practices (deliberately excluding "
    "hospitals) along with the HR teams who hire on their behalf, and the company served a range of medical "
    "specialties. That ICP and specialty mix directly informed which tools, page types, and positioning I built "
    "into the site."
)

subhead("The architecture I designed and the reasoning behind each part")
body(
    "After competitor research, including studying how **Indeed** structured its hiring content as a "
    "keyword-rich hub, I designed the full framework and produced the content for it:"
)
bullet(
    "**Homepage** rebuilt to rank for the top keywords in the medical virtual staffing industry. I included a "
    "**custom interactive cost calculator** I spec'd and built (using ChatGPT) and had the developer implement, "
    "so prospects could calculate their staffing costs and see exactly what we offered. The page flowed from a "
    "hero section (the core pitch) into announcements, a trust bar (partners), a “why work with us versus "
    "other agencies” section, and **role cards** that linked through to the contact page. Apart from internal "
    "linking, virtually every CTA routed to the contact page, and I also added a deliberately simple inline form "
    "so prospects could convert immediately."
)
bullet(
    "**Role / service pages (12)** built to reflect the actual roles being hired for, the pricing, and "
    "everything the SDRs promise on calls. This mattered because virtual-staffing pricing is two-sided: the "
    "client-facing rate differs from what the VA is paid (for example, a role billed to the client around "
    "**$17** might pay the VA roughly **$7 to $8**). The roles were specialized, including medical nurse, dental "
    "nurse, and dental admin, with distinct senior and junior tiers, so I worked with the marketing team to get "
    "the correct facts and figures before writing."
)
bullet(
    "**Specialties hub plus ~18 specialty pages**, created after competitor research to extend our reach across "
    "the individual specialties the company served."
)
bullet(
    "**Capabilities hub plus ~16 capability pages**, built to capture searches for specific functions rather "
    "than only “virtual medical assistant.” Someone may be searching for an insurance verification "
    "specialist or a billing specialist without using the umbrella term, and I wanted us to show up for those, "
    "building topical keyword clusters and authority around each."
)
bullet(
    "**How-to-Hire section (~43 pages)**, the job-description and hiring-focused pages and the last major piece "
    "I built before leaving. The idea came from the Indeed audit: this hub targets both doctors hiring directly "
    "and the HR teams of practices actively looking to partner or hire VAs, by ranking for the keywords they use "
    "when writing job descriptions."
)
bullet(
    "**Blog hub**, intentionally my primary traffic-building engine, since the existing blog was underdeveloped "
    "and generic. I curated quality content here to drive qualified traffic to the site."
)
bullet("**Announcements** for internal announcements and company promos.")
bullet(
    "**Newsletters / press releases** as part of a backlink-building strategy: every press piece we published "
    "was also posted on-site with the relevant links so we captured the link juice, complemented by guest "
    "posting and similar outreach."
)
bullet(
    "**FAQ hub**, a dedicated hub built to serve both SEO and LLMs as a citation source. Each page's FAQs linked "
    "back to this hub, creating internal-linking depth, and I added **FAQ schema** so the content could be "
    "surfaced as concise, direct answers."
)
bullet(
    "**Location pages**, added to capture US traffic beyond just California (the company served the whole "
    "country); these were in progress and largely finished with UI/UX at handoff."
)
bullet(
    "**Supporting pages**: About Us with a team page, Contact (I drafted the brief), Thank You, plus standard "
    "Terms of Service, Cookie Policy, and other legal pages."
)

subhead("Why I sequenced content before design")
body(
    "At the time I joined there was no UI/UX designer, so I handled that thinking myself initially, writing the "
    "copy with UI/UX in mind and drawing on my own design background. I was particular about avoiding "
    "text-heaviness; I wanted a clean, flowing site, so I made sure the eventual design fit what I considered "
    "best practice. Some companies design first, but given the staffing situation, content first was the right "
    "call here. Once briefs were ready, UI/UX designed the pages, sent them back to me, and then they went into "
    "development."
)

subhead("Methodology across every page")
body(
    "For each page I conducted dedicated keyword research, reviewed competitors and what we were already ranking "
    "for, studied the SERP, and built **topic clusters** rather than chasing single keywords, the goal being "
    "topical authority strong enough to outrank competitors. I reviewed every piece of copy to confirm it "
    "aligned with what we actually promised, cross-checked it against Google's **E-E-A-T** criteria (using "
    "ChatGPT as a check), and optimized for **LLM-friendliness** with concise, direct answers and schema."
)

subhead("Scale and timeline")
body(
    "In total this was roughly **100 pages** (12 role pages, ~16 capability pages, ~18 specialty pages, ~43 "
    "how-to-hire pages, plus the hubs, homepage, and supporting/legal pages). I worked in **30/60/90-day "
    "buckets**. The first batch, the 12 role pages, ~10 structural/hub pages, and an initial set of capability "
    "and specialty pages, took about a month from content production through design review and into development. "
    "After that I sustained a cadence of roughly **5 to 10 pages per week**. All of this ran alongside my core "
    "SEO specialist duties: writing blogs, handling on-page and off-page optimization, running audits, and "
    "automating our content-writing system (including a six-month blog plan I could generate briefs from on "
    "demand, populate with live information, generate, then edit to read as human-written)."
)

subhead("Results")
body(
    "It took about three to six months to begin seeing meaningful growth and roughly **8 months to a year** to "
    "hit our stride. Organic traffic grew from an initial **12 to 27 mostly-direct visits per month** to about "
    "**8,000 visits per month**. The lever was using blog content to pull in additional ranking keywords and "
    "internally linking those posts to relevant commercial pages; I then re-optimized pages to rank higher for "
    "commercial keywords rather than purely informational ones. The pages were also continuously updated to keep "
    "pace with Google updates, since those routinely shift rankings. One concrete challenge: some content "
    "initially ranked in India rather than the US, which I corrected by updating the on-page signals, adding "
    "US-specific terminology and local sources, so it ranked specifically in the US. All of this is visible in "
    "Ahrefs, and I am happy to share access or screenshots."
)

# ---------------- Q2 ----------------
question_block(
    "Q2",
    "Have you used AI as part of a repeatable content production process? Describe the workflow, including "
    "prompting, editing, and QA steps, and the approximate number of pages produced.",
)

body(
    "Yes. I built a repeatable **content intelligence system** at Med Virtual by rewriting my own editorial "
    "workflow into a documented, reusable process, not a series of one-off prompts. I used it most heavily to "
    "produce the **~43 how-to-hire / job-description pages** (all generated in roughly a day or two), and I also "
    "used it to automate the blog program, the press releases, and the newsletters. I am happy to share the "
    "actual pages and the underlying system structure if useful."
)

subhead("Tools and infrastructure")
body(
    "The core engine was **Claude / Claude Code**, run on my personal account (about $200 per month) because the "
    "agency's account was on a lower, restrictive tier where you would constantly hit limits. Alongside Claude I "
    "used **DeepSeek** and **ChatGPT** as interpreters and refiners, **Surfer SEO** as a prompt enhancer, and "
    "**Notion** (free tier) as the knowledge database. I did not need GitHub or Supabase for this; since I was "
    "not on Next.js, Notion was sufficient as the single source of truth."
)

subhead("What lived in the database")
body(
    "Notion stored everything the system reasoned from: Med Virtual's **brand guidelines**, my **rules for what "
    "not to include** (banned phrases and sentence structures), and the definition of what qualified as a "
    "“good” blog. The system also held the defined **phases** of my process, phase 0 through ~18, so "
    "the workflow ran consistently each time. I can share what that phase structure looks like."
)

subhead("How I trained it")
body(
    "Rather than writing long static prompts, I trained it **conversationally**, passing my knowledge to it and "
    "then converting that into a repeatable, reusable skill. I first went online, found content I considered "
    "strong, and taught it what “good” looked like: which sections a piece needs, when to use an "
    "infographic versus a table, how many sentences belong in a paragraph, how to open a piece, when to include "
    "a key-takeaways section, how to approach a content piece, when and how to reference Med Virtual, and how to "
    "suggest images. Conversational training consistently worked better for me than rigid prompting."
)

subhead("Prompting workflow")
body(
    "To make sure the model understood my intent precisely, I used a refinement chain: I would brief "
    "conversationally, run it through **Surfer SEO's prompt enhancer** to translate my intent into LLM-ready "
    "phrasing, then refine it in **DeepSeek** (which I found the best interpreter of my intent) and **ChatGPT**. "
    "Once the refined prompt genuinely represented my thinking and process, I would paste it into **Claude "
    "Code** to generate."
)

subhead("Data integrations (MCPs and APIs)")
body(
    "I connected **Ahrefs** and **Semrush** via MCP so I could make API calls, and added **DataForSEO** for "
    "AI/LLM data. For keyword research I deliberately looked for **confluence across all three platforms**, "
    "since each measures volume and metrics differently. A keyword had to show solid search volume on all three, "
    "healthy CPC, US dominance, and the right intent (commercial or informational, depending on the piece). I "
    "then checked it against what was currently ranking in the SERP and against the content cluster it belonged "
    "to in my existing plan (originally a quarterly plan I expanded into a six-month plan). I also connected "
    "**Google Search Console** and **Google Analytics** API keys plus **Microsoft Clarity** (initially via "
    "Claude controlling the browser in Chrome, before I moved fully to Claude Code) so the system could pull "
    "recent queries related to what I was writing, find pages with **high impressions but low clicks** to "
    "optimize, and identify top-performing pages, confirmed across GSC (primary), Ahrefs, and Semrush, so we "
    "knew which pages not to touch and what to learn from them. **Canva** was connected for imagery, but since "
    "AI image output was not always accurate, I often had it produce infographics in **HTML** or hand me a "
    "detailed spec I would finalize in Canva myself, depending on complexity."
)

subhead("Content generation")
body(
    "Because I had already built detailed generation instructions with specific QA and clarity checks baked in, "
    "the system could weave it all together and produce a piece that was pushed to **Notion** and exported to "
    "**Google Docs**."
)

subhead("Editing and QA, the human-in-the-loop step")
body(
    "Every piece ran through a defined QA and clarity check covering **E-E-A-T**, **LLM-friendliness**, **FAQ "
    "schema**, and **brand-guideline alignment**. Critically, I kept a human in the loop on every piece: I had "
    "the option to connect directly to Webflow but chose not to, because I preferred to take the finished draft "
    "out, **read it, verify it, identify any needed optimizations, and edit it to read as genuinely "
    "human-written** before publishing it into Webflow myself. Google rewards content written for humans, not AI "
    "filler, so nothing went live unverified or irrelevant."
)

subhead("Approximate volume produced")
body(
    "Through this process I produced roughly **100+ pieces** in total: the ~43 how-to-hire pages, the ongoing "
    "blog program (the posts that drove the bulk of our traffic growth), roughly **12 to 14 newsletters**, and "
    "the press releases."
)

# Footer
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Joseph  •  SEO Content & Strategy  •  Screening Questionnaire Response")
fr.font.size = Pt(8.5)
fr.font.color.rgb = LIGHT

doc.save("/home/user/Workout-Coach/SEO_Content_Questionnaire_Joseph.docx")
print("saved")
